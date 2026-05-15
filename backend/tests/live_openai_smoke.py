from __future__ import annotations

import json
import os
import sys
from copy import deepcopy
from io import BytesIO
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
SMOKE_DIR = ROOT / ".pytest-tmp" / "live-openai-smoke"
SMOKE_DIR.mkdir(parents=True, exist_ok=True)
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

os.environ.setdefault("DOCMAP_LOCAL_FLOW_STORE", ".pytest-tmp/live-openai-smoke/flows.json")
os.environ.setdefault(
    "DOCMAP_LOCAL_AI_DRAFT_SESSION_STORE",
    ".pytest-tmp/live-openai-smoke/ai_draft_sessions.json",
)
os.environ.setdefault("openai_fast_model", os.getenv("DOCMAP_LIVE_SMOKE_MODEL", "gpt-5.4-mini"))
os.environ.setdefault("openai_balanced_model", os.getenv("DOCMAP_LIVE_SMOKE_MODEL", "gpt-5.4-mini"))
os.environ.setdefault("openai_default_model", os.getenv("DOCMAP_LIVE_SMOKE_MODEL", "gpt-5.4-mini"))

from config import reset_request_settings, set_request_settings  # noqa: E402
from documents.ingestion import ingest_supported_document  # noqa: E402
from export.workspace_graph import build_workspace_graph  # noqa: E402
from ai_helpers import (  # noqa: E402
    accept_ai_draft_revision,
    discard_ai_draft_session,
    generate_ai_draft_session_with_provider,
    revise_ai_draft_session_with_provider,
)
from app import (  # noqa: E402
    _append_accepted_graph_to_flow_snapshot,
    _persist_flow_snapshot,
    export_workspace_json,
    export_workspace_markdown,
    get_workspace_graph_or_404,
    local_create_flow,
    save_ai_draft_session,
)


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _small_docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Breaker Panel Label Review", level=1)
    document.add_paragraph(
        "Panel A feeds lighting circuits in the north hall. Technicians must inspect breaker "
        "labels before export and preserve the panel identifier in notes."
    )
    document.add_paragraph(
        "If a circuit label is missing or uncertain, mark the generated node as needs_review "
        "and attach a review note rather than inventing a circuit name."
    )
    document.add_paragraph(
        "Source-backed draft items should cite this document and use concise operational wording."
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _react_flow_snapshot() -> dict:
    return {
        "nodes": [
            {
                "id": "root-panel-review",
                "type": "custom",
                "position": {"x": 80, "y": 120},
                "data": {
                    "title": "Panel Review",
                    "summ": "Draft source-grounded review steps for breaker panel labels.",
                    "node_type": "concept",
                },
            }
        ],
        "edges": [],
        "viewport": {"x": 0, "y": 0, "zoom": 1},
    }


def _source_chunks_from_docx() -> list[dict]:
    ingested = ingest_supported_document("live-openai-smoke.docx", _small_docx_bytes(), version=1)
    chunks = []
    for chunk in ingested["document_chunks"][:2]:
        quote = str(chunk.get("text") or "")[:180]
        chunks.append(
            {
                **chunk,
                "source_ref": {
                    "document_id": chunk["document_id"],
                    "chunk_id": chunk["id"],
                    "page": chunk.get("page"),
                    "section": chunk.get("heading") or "Breaker Panel Label Review",
                    "quote_snippet": quote,
                    "confidence": "medium",
                },
            }
        )
    return chunks


def _selected_node_item_ids(session: dict) -> list[str]:
    latest = session["revisions"][-1]
    ids = []
    for item in latest.get("draft_items", []):
        metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
        if item.get("item_type") == "node" or metadata.get("node_id"):
            ids.append(item["id"])
    if ids:
        return ids[:1]
    nodes = latest.get("draft_nodes", [])
    return [nodes[0]["id"]] if nodes else []


def _node_terms(nodes: list[dict]) -> set[str]:
    terms = set()
    for node in nodes:
        text = f"{node.get('title', '')} {node.get('summary', '')}".lower()
        for term in ("panel", "breaker", "label", "review", "circuit"):
            if term in text:
                terms.add(term)
    return terms


def run() -> dict:
    api_key = os.getenv("openai_api_key") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "MISSING_OPENAI_API_KEY: set openai_api_key or OPENAI_API_KEY to run the live Responses smoke."
        )

    token = set_request_settings({"openai_api_key": api_key})
    try:
        flow = local_create_flow(
            {
                "flow_name": "Live OpenAI Smoke",
                "flow_type": "manual",
                "summary": "Disposable live Responses smoke workspace.",
                "flow_json": json.dumps(_react_flow_snapshot()),
            }
        )
        flow_id = str(flow["_id"])
        graph = build_workspace_graph(flow)
        source_chunks = _source_chunks_from_docx()
        before_graph = deepcopy(graph)

        session = generate_ai_draft_session_with_provider(
            graph,
            workspace_id=flow_id,
            prompt=(
                "Quick source-grounded draft: create two concise graph nodes about breaker "
                "panel label inspection. Cite only the provided source refs."
            ),
            scope={"type": "node", "node_id": "root-panel-review"},
            role="Ask AI Live Smoke",
            model_policy="speed",
            model=os.getenv("DOCMAP_LIVE_SMOKE_MODEL", "gpt-5.4-mini"),
            desired_outputs=["mind_map"],
            source_chunks=source_chunks,
        )
        save_ai_draft_session(session)

        first_revision = session["revisions"][-1]
        _assert(session.get("selected_model"), "draft session did not record selected_model")
        _assert(session.get("model_reason"), "draft session did not record model_reason")
        _assert(first_revision.get("draft_nodes"), "live draft did not produce graph draft nodes")
        _assert(session.get("source_refs"), "live draft did not preserve source refs")
        _assert(_node_terms(first_revision["draft_nodes"]), "live draft nodes did not reflect source topic")

        revised = revise_ai_draft_session_with_provider(
            session,
            graph,
            prompt=(
                "Follow-up revise: include one short item that flags missing or uncertain labels "
                "for needs_review without inventing circuit names."
            ),
            model_policy="speed",
            model=os.getenv("DOCMAP_LIVE_SMOKE_MODEL", "gpt-5.4-mini"),
            desired_outputs=["mind_map", "review_annotations"],
            source_chunks=source_chunks,
        )
        save_ai_draft_session(revised)
        _assert(len(revised["revisions"]) == len(session["revisions"]) + 1, "follow-up did not add a draft revision")
        _assert(revised["revisions"][-1]["revision_id"] != first_revision["revision_id"], "revision id did not change")

        selected_ids = _selected_node_item_ids(revised)
        _assert(selected_ids, "no selectable draft node items found")
        accepted_graph, accepted_session, accept_result = accept_ai_draft_revision(
            graph,
            revised,
            accept_mode="selected",
            selected_item_ids=selected_ids,
            accepted_by="live-openai-smoke",
        )
        save_ai_draft_session(accepted_session)
        _assert(accept_result.get("accepted_node_ids"), "selected accept did not accept any nodes")
        _assert(accept_result.get("graph_revision_id"), "accept did not create a graph revision id")
        _assert(accept_result.get("undo", {}).get("before_graph") == before_graph, "accept missing undoable before graph")

        accepted_nodes = [
            node
            for node in accepted_graph.get("nodes", [])
            if node.get("id") in set(accept_result["accepted_node_ids"])
        ]
        _assert(accepted_nodes, "accepted nodes not present in canonical graph")
        _assert(
            all(node.get("source_refs") or node.get("status") == "needs_review" for node in accepted_nodes),
            "accepted nodes are neither cited nor marked needs_review",
        )

        snapshot = _append_accepted_graph_to_flow_snapshot(flow, accept_result, accepted_graph)
        _persist_flow_snapshot(flow_id, snapshot)
        reloaded_graph = get_workspace_graph_or_404(flow_id)
        reloaded_nodes = [
            node
            for node in reloaded_graph.get("nodes", [])
            if node.get("id") in set(accept_result["accepted_node_ids"])
        ]
        _assert(len(reloaded_nodes) == len(accepted_nodes), "save/reload lost accepted nodes")
        _assert(
            all(node.get("source_refs") or node.get("status") == "needs_review" for node in reloaded_nodes),
            "save/reload lost source refs or needs_review state",
        )

        discarded = discard_ai_draft_session(revised, discarded_by="live-openai-smoke")
        _assert(
            build_workspace_graph(flow) == before_graph,
            "discarding/canceling a draft mutated the pre-accept graph",
        )
        _assert(discarded.get("status") == "discarded", "discard did not mark draft as discarded")

        exported_json = export_workspace_json(flow_id)
        exported_markdown = export_workspace_markdown(flow_id).body.decode("utf-8")
        _assert(exported_json.get("nodes"), "JSON export missing nodes")
        _assert("Live OpenAI Smoke" in exported_markdown, "Markdown export missing workspace title")

        result = {
            "flow_id": flow_id,
            "model": accepted_session.get("selected_model"),
            "model_reason": accepted_session.get("model_reason"),
            "draft_revision_count": len(accepted_session.get("revisions", [])),
            "accepted_node_ids": accept_result.get("accepted_node_ids", []),
            "source_refs_survived_reload": all(bool(node.get("source_refs")) for node in reloaded_nodes),
            "needs_review_survived_reload": any(node.get("status") == "needs_review" for node in reloaded_nodes),
            "graph_revision_id": accept_result.get("graph_revision_id"),
            "undo_kind": accept_result.get("undo", {}).get("kind"),
            "export_json_node_count": len(exported_json.get("nodes", [])),
            "export_markdown_chars": len(exported_markdown),
            "discarded_status": discarded.get("status"),
        }
        (SMOKE_DIR / "result.json").write_text(json.dumps(result, indent=2), encoding="utf-8")
        return result
    finally:
        reset_request_settings(token)


if __name__ == "__main__":
    try:
        print(json.dumps(run(), indent=2))
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        raise
