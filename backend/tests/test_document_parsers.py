import sys
import base64
from io import BytesIO
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from documents.ingestion import (
    DocumentIngestionError,
    chunk_source_segments,
    extract_source_segments,
    ingest_supported_document,
    source_segments_from_docx,
    source_segments_from_pdf,
    source_segments_from_text,
)


FIXTURE_DIR = Path(__file__).parent / "fixtures"
PDF_FIXTURE_B64 = (
    "JVBERi0xLjMKJZOMi54gUmVwb3J0TGFiIEdlbmVyYXRlZCBQREYgZG9jdW1lbnQg"
    "aHR0cDovL3d3dy5yZXBvcnRsYWIuY29tCjEgMCBvYmoKPDwKL0YxIDIgMCBSCj4+"
    "CmVuZG9iagoyIDAgb2JqCjw8Ci9CYXNlRm9udCAvSGVsdmV0aWNhIC9FbmNvZGlu"
    "ZyAvV2luQW5zaUVuY29kaW5nIC9OYW1lIC9GMSAvU3VidHlwZSAvVHlwZTEgL1R5"
    "cGUgL0ZvbnQKPj4KZW5kb2JqCjMgMCBvYmoKPDwKL0NvbnRlbnRzIDggMCBSIC9N"
    "ZWRpYUJveCBbIDAgMCA1OTUuMjc1NiA4NDEuODg5OCBdIC9QYXJlbnQgNyAwIFIg"
    "L1Jlc291cmNlcyA8PAovRm9udCAxIDAgUiAvUHJvY1NldCBbIC9QREYgL1RleHQg"
    "L0ltYWdlQiAvSW1hZ2VDIC9JbWFnZUkgXQo+PiAvUm90YXRlIDAgL1RyYW5zIDw8"
    "Cgo+PiAKICAvVHlwZSAvUGFnZQo+PgplbmRvYmoKNCAwIG9iago8PAovQ29udGVu"
    "dHMgOSAwIFIgL01lZGlhQm94IFsgMCAwIDU5NS4yNzU2IDg0MS44ODk4IF0gL1Bh"
    "cmVudCA3IDAgUiAvUmVzb3VyY2VzIDw8Ci9Gb250IDEgMCBSIC9Qcm9jU2V0IFsg"
    "L1BERiAvVGV4dCAvSW1hZ2VCIC9JbWFnZUMgL0ltYWdlSSBdCj4+IC9Sb3RhdGUg"
    "MCAvVHJhbnMgPDwKCj4+IAogIC9UeXBlIC9QYWdlCj4+CmVuZG9iago1IDAgb2Jq"
    "Cjw8Ci9QYWdlTW9kZSAvVXNlTm9uZSAvUGFnZXMgNyAwIFIgL1R5cGUgL0NhdGFs"
    "b2cKPj4KZW5kb2JqCjYgMCBvYmoKPDwKL0F1dGhvciAoYW5vbnltb3VzKSAvQ3Jl"
    "YXRpb25EYXRlIChEOjIwMjYwNTE0MTMxMjQwLTA0JzAwJykgL0NyZWF0b3IgKFJl"
    "cG9ydExhYiBQREYgTGlicmFyeSAtIHd3dy5yZXBvcnRsYWIuY29tKSAvS2V5d29y"
    "ZHMgKCkgL01vZERhdGUgKEQ6MjAyNjA1MTQxMzEyNDAtMDQnMDAnKSAvUHJvZHVj"
    "ZXIgKFJlcG9ydExhYiBQREYgTGlicmFyeSAtIHd3dy5yZXBvcnRsYWIuY29tKSAK"
    "ICAvU3ViamVjdCAodW5zcGVjaWZpZWQpIC9UaXRsZSAodW50aXRsZWQpIC9UcmFw"
    "cGVkIC9GYWxzZQo+PgplbmRvYmoKNyAwIG9iago8PAovQ291bnQgMiAvS2lkcyBb"
    "IDMgMCBSIDQgMCBSIF0gL1R5cGUgL1BhZ2VzCj4+CmVuZG9iago4IDAgb2JqCjw8"
    "Ci9GaWx0ZXIgWyAvQVNDSUk4NURlY29kZSAvRmxhdGVEZWNvZGUgXSAvTGVuZ3Ro"
    "IDE2MAo+PgpzdHJlYW0KR2FyVzBZbVM/NSY0SERDYEtcQGRoTS1RbyYwWUhFcmtR"
    "KFJAVXNOZGhXPWMjTi5sdTAwWj5fYSNkL3RJLjItWjAqNm5cJV8wOGAtRlBtT1FZ"
    "IkRMSWwqMmdgaSFJJFA8OlhmPURuKnMmQkZURUJwTzwxUCFDS1E2Um1IMFJlaFAo"
    "XyE6ci0sLzNdQG5vJEZWMFZXUiVLZUptTCEubXN+PmVuZHN0cmVhbQplbmRvYmoK"
    "OSAwIG9iago8PAovRmlsdGVyIFsgL0FTQ0lJODVEZWNvZGUgL0ZsYXRlRGVjb2Rl"
    "IF0gL0xlbmd0aCAxNTIKPj4Kc3RyZWFtCkdhcFFoMEU9RiwwVVxIM1RccE5ZVF5R"
    "S2s/dGM+SVAsO1cjVTFeMjNpaFBFTV8/Q1c0S0lTaTwhWzdgI09CX3NLNVkkTyRl"
    "P1s8SkVfVjdMLzs0V2UvMU0pTy5SIUJcJFc8MF0nVSI+Tyc9MVZxJ18ybFgxXDkk"
    "MFAjQmQiTWMoOWYuZ2tqKjkoXDM9OztfS1VRWkZMT34+ZW5kc3RyZWFtCmVuZG9i"
    "agp4cmVmCjAgMTAKMDAwMDAwMDAwMCA2NTUzNSBmIAowMDAwMDAwMDczIDAwMDAw"
    "IG4gCjAwMDAwMDAxMDQgMDAwMDAgbiAKMDAwMDAwMDIxMSAwMDAwMCBuIAowMDAw"
    "MDAwNDE0IDAwMDAwIG4gCjAwMDAwMDA2MTcgMDAwMDAgbiAKMDAwMDAwMDY4NSAw"
    "MDAwMCBuIAowMDAwMDAwOTgxIDAwMDAwIG4gCjAwMDAwMDEwNDYgMDAwMDAgbiAK"
    "MDAwMDAwMTI5NiAwMDAwMCBuIAp0cmFpbGVyCjw8Ci9JRCAKWzxhMTBlYjQ0YmM5"
    "NGJmM2NmN2Y5N2JlYzJmYjZlZjY5OD48YTEwZWI0NGJjOTRiZjNjZjdmOTdiZWMy"
    "ZmI2ZWY2OTg+XQolIFJlcG9ydExhYiBnZW5lcmF0ZWQgUERGIGRvY3VtZW50IC0t"
    "IGRpZ2VzdCAoaHR0cDovL3d3dy5yZXBvcnRsYWIuY29tKQoKL0luZm8gNiAwIFIK"
    "L1Jvb3QgNSAwIFIKL1NpemUgMTAKPj4Kc3RhcnR4cmVmCjE1MzgKJSVFT0YK"
)


def test_markdown_fixture_extracts_headings_and_offsets():
    text = (FIXTURE_DIR / "sample.md").read_text(encoding="utf-8")

    segments = source_segments_from_text(text, "md")
    chunks = chunk_source_segments(segments, "src_markdown_v1")

    assert segments[0]["heading"] == "Electrical Scope"
    assert segments[-1]["heading"] == "QA Review"
    assert chunks[0]["document_id"] == "src_markdown_v1"
    assert chunks[0]["start_char"] == segments[0]["start_char"]


def test_txt_fixture_extracts_character_locations():
    text = (FIXTURE_DIR / "sample.txt").read_text(encoding="utf-8")

    segments = source_segments_from_text(text, "txt")

    assert segments[0]["heading"] == "Electrical Scope"
    assert segments[0]["start_char"] == 0
    assert segments[-1]["end_char"] <= len(text)


def test_docx_parser_extracts_paragraph_heading_context():
    file_bytes = _build_docx_fixture()

    segments = source_segments_from_docx(file_bytes)

    assert segments[0]["heading"] == "Electrical Scope"
    assert segments[1]["heading"] == "Electrical Scope"
    assert segments[1]["text"] == "Install conduit and label panel schedules."
    assert segments[1]["start_char"] > segments[0]["start_char"]


def test_pdf_parser_extracts_page_locations():
    file_bytes = _build_pdf_fixture()

    segments = source_segments_from_pdf(file_bytes)

    assert segments
    assert segments[0]["page"] == 1
    assert "Electrical Scope" in segments[0]["text"]


def test_extract_source_segments_dispatches_by_supported_type():
    md_bytes = (FIXTURE_DIR / "sample.md").read_bytes()

    segments = extract_source_segments(md_bytes, "md")

    assert segments[0]["heading"] == "Electrical Scope"


def test_mvp_upload_extract_chunk_flow_for_all_supported_fixtures():
    fixtures = {
        "pdf": ("sample.pdf", _build_pdf_fixture()),
        "docx": ("sample.docx", _build_docx_fixture()),
        "md": ("sample.md", (FIXTURE_DIR / "sample.md").read_bytes()),
        "txt": ("sample.txt", (FIXTURE_DIR / "sample.txt").read_bytes()),
    }

    for extension, (filename, file_bytes) in fixtures.items():
        result = ingest_supported_document(filename, file_bytes)

        assert result["source_document"]["type"] == extension
        assert result["source_document"]["filename"] == filename
        assert result["source_document"]["file_hash"]
        assert result["source_segments"]
        assert result["document_chunks"]
        assert result["document_chunks"][0]["document_id"] == result["source_document"]["id"]
        assert result["document_chunks"][0]["id"].startswith("chk_")


def test_mvp_upload_extract_chunk_flow_rejects_non_primary_formats():
    try:
        ingest_supported_document("legacy-demo.html", b"<h1>Not MVP</h1>")
    except DocumentIngestionError as exc:
        assert "Unsupported file type" in str(exc)
    else:
        raise AssertionError("Expected unsupported HTML upload to be rejected.")


def _build_docx_fixture() -> bytes:
    from docx import Document

    stream = BytesIO()
    document = Document()
    document.add_heading("Electrical Scope", level=1)
    document.add_paragraph("Install conduit and label panel schedules.")
    document.add_heading("QA Review", level=2)
    document.add_paragraph("Inspect breaker naming before export.")
    document.save(stream)
    return stream.getvalue()


def _build_pdf_fixture() -> bytes:
    return base64.b64decode(PDF_FIXTURE_B64)
