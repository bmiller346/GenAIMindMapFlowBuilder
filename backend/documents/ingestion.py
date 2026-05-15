from __future__ import annotations

import hashlib
import os
import re
from dataclasses import asdict, dataclass
from pathlib import PurePath
from zipfile import BadZipFile

ALLOWED_DOCUMENT_EXTENSIONS = frozenset({"pdf", "docx", "md", "txt"})
ALLOWED_AI_INTAKE_EXTENSIONS = frozenset({"pdf", "docx", "md", "txt", "pptx", "html"})
SOURCE_TRACEABLE_DOCUMENT_LABEL = "PDF, DOCX, Markdown, or TXT"
DEFAULT_MAX_UPLOAD_BYTES = 25 * 1024 * 1024
MAX_CHUNK_CHARS = 2200
MIN_CHUNK_CHARS = 400


class DocumentIngestionError(ValueError):
    """Raised when an uploaded source document cannot enter the TraceSpace pipeline."""


@dataclass(slots=True)
class SourceDocument:
    id: str
    filename: str
    original_filename: str
    type: str
    file_hash: str
    size: int
    version: int
    status: str = "uploaded"


@dataclass(slots=True)
class DocumentChunk:
    id: str
    document_id: str
    index: int
    text: str
    page: int | None = None
    heading: str | None = None
    start_char: int = 0
    end_char: int = 0


@dataclass(slots=True)
class SourceSegment:
    text: str
    page: int | None = None
    heading: str | None = None
    start_char: int = 0
    end_char: int = 0


def configured_max_upload_bytes() -> int:
    raw_value = os.getenv("DOCMAP_MAX_UPLOAD_BYTES")
    if not raw_value:
        return DEFAULT_MAX_UPLOAD_BYTES

    try:
        return int(raw_value)
    except ValueError as exc:
        raise DocumentIngestionError("DOCMAP_MAX_UPLOAD_BYTES must be an integer.") from exc


def sanitize_filename(filename: str) -> str:
    original = PurePath(filename or "document").name
    stem, extension = os.path.splitext(original)
    extension = extension.lower().lstrip(".")

    safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "-", stem).strip("._-")
    if not safe_stem:
        safe_stem = "document"

    if extension:
        return f"{safe_stem[:120]}.{extension}"
    return safe_stem[:120]


def file_sha256(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def validate_upload_bytes(filename: str, file_bytes: bytes) -> dict:
    return validate_file_upload_bytes(
        filename,
        file_bytes,
        ALLOWED_DOCUMENT_EXTENSIONS,
        unsupported_message=(
            "This source cannot enter the source-traceable document pipeline. "
            f"Upload {SOURCE_TRACEABLE_DOCUMENT_LABEL} for document-section citations, "
            "or use the matching AI intake option for non-document sources."
        ),
    )


def validate_ai_intake_bytes(filename: str, file_bytes: bytes) -> dict:
    return validate_file_upload_bytes(filename, file_bytes, ALLOWED_AI_INTAKE_EXTENSIONS)


def validate_file_upload_bytes(
    filename: str,
    file_bytes: bytes,
    allowed_extensions: frozenset[str],
    unsupported_message: str = "",
) -> dict:
    sanitized_filename = sanitize_filename(filename)
    extension = sanitized_filename.rsplit(".", 1)[-1].lower() if "." in sanitized_filename else ""

    if extension not in allowed_extensions:
        allowed = ", ".join(sorted(allowed_extensions))
        guidance = unsupported_message or f"Allowed types: {allowed}."
        raise DocumentIngestionError(
            f"Unsupported file type '{extension}'. {guidance}"
        )

    if not file_bytes:
        raise DocumentIngestionError("The uploaded file is empty.")

    max_size = configured_max_upload_bytes()
    if len(file_bytes) > max_size:
        max_mb = max_size / (1024 * 1024)
        raise DocumentIngestionError(f"File exceeds the {max_mb:.1f} MB upload limit.")

    return {
        "filename": sanitized_filename,
        "original_filename": filename or sanitized_filename,
        "extension": extension,
        "size": len(file_bytes),
        "file_hash": file_sha256(file_bytes),
    }


def build_source_document(filename: str, file_bytes: bytes, version: int = 1) -> dict:
    upload = validate_upload_bytes(filename, file_bytes)
    return source_document_from_upload(upload, version)


def build_ai_intake_source_document(filename: str, file_bytes: bytes, version: int = 1) -> dict:
    upload = validate_ai_intake_bytes(filename, file_bytes)
    return source_document_from_upload(upload, version)


def source_document_from_upload(upload: dict, version: int = 1) -> dict:
    source_document = SourceDocument(
        id=f"src_{upload['file_hash'][:16]}_v{version}",
        filename=upload["filename"],
        original_filename=upload["original_filename"],
        type=upload["extension"],
        file_hash=upload["file_hash"],
        size=upload["size"],
        version=version,
    )
    return asdict(source_document)


def normalize_relative_source_path(relative_path: str | None, fallback_filename: str = "") -> str:
    raw_path = str(relative_path or "").replace("\\", "/").strip()
    raw_path = re.sub(r"^[A-Za-z]:/+", "", raw_path).lstrip("/")
    parts = [
        sanitize_filename(part)
        for part in raw_path.split("/")
        if part and part not in {".", ".."}
    ]

    if not parts and fallback_filename:
        parts = [sanitize_filename(fallback_filename)]

    return "/".join(part for part in parts if part)


def build_source_set_metadata(
    relative_paths: list[str],
    *,
    source_set_id: str | None = None,
    label: str | None = None,
) -> dict:
    normalized_paths = [
        normalize_relative_source_path(path)
        for path in relative_paths
        if normalize_relative_source_path(path)
    ]
    digest_input = "|".join(sorted(normalized_paths)) or "source-set"
    digest = hashlib.sha256(digest_input.encode("utf-8")).hexdigest()[:12]
    root_folder = _common_source_root(normalized_paths)
    clean_id = _stable_source_set_token(source_set_id or f"source_set_{digest}")

    return {
        "id": clean_id,
        "label": clean_source_set_label(label) or root_folder or "Uploaded source set",
        "root_folder": root_folder,
        "source_count": len(normalized_paths),
        "upload_mode": "native_folder_upload",
        "native_folder_upload": True,
    }


def source_document_with_source_set_metadata(
    source_document: dict,
    *,
    relative_path: str,
    source_set: dict,
) -> dict:
    document = dict(source_document)
    normalized_path = normalize_relative_source_path(
        relative_path,
        fallback_filename=document.get("filename", ""),
    )
    folder = normalized_path.rsplit("/", 1)[0] if "/" in normalized_path else ""
    source_set_record = dict(source_set)

    document.update(
        {
            "relative_path": normalized_path,
            "path": normalized_path,
            "folder": folder,
            "source_set_id": source_set_record.get("id", ""),
            "source_set": source_set_record,
        }
    )
    return document


def ingest_supported_document(filename: str, file_bytes: bytes, version: int = 1) -> dict:
    upload = validate_upload_bytes(filename, file_bytes)
    source_document = build_source_document(upload["filename"], file_bytes, version=version)
    source_segments = extract_source_segments(file_bytes, source_document["type"])
    document_chunks = chunk_source_segments(source_segments, source_document["id"])

    if not document_chunks:
        raise DocumentIngestionError("Document did not produce any source-aware sections.")

    return {
        "upload": upload,
        "file_bytes": file_bytes,
        "source_document": source_document,
        "source_segments": source_segments,
        "document_chunks": document_chunks,
    }


def deterministic_chunk_id(document_id: str, index: int, text: str) -> str:
    digest = hashlib.sha256(f"{document_id}:{index}:{text}".encode("utf-8")).hexdigest()
    return f"chk_{digest[:20]}"


def chunk_text(text: str, document_id: str, *, page: int | None = None) -> list[dict]:
    return chunk_source_segments(
        [SourceSegment(text=text, page=page, start_char=0, end_char=len(text))],
        document_id,
    )


def chunk_source_segments(segments: list[SourceSegment] | list[dict], document_id: str) -> list[dict]:
    source_segments = [_coerce_segment(segment) for segment in segments]
    chunks: list[DocumentChunk] = []

    for segment in source_segments:
        normalized = _normalize_text(segment.text)
        if not normalized:
            continue

        offset = segment.start_char
        current_heading = segment.heading
        current_text = ""
        current_start = offset
        cursor = offset

        for block in _iter_blocks(normalized):
            block_text = block["text"]
            block_start = offset + block["start"]
            heading = _heading_from_block(block_text)
            if heading:
                current_heading = heading

            if current_text and len(current_text) + len(block_text) + 2 > MAX_CHUNK_CHARS:
                _append_chunk(
                    chunks,
                    document_id,
                    current_text,
                    current_start,
                    cursor,
                    segment.page,
                    current_heading,
                )
                current_text = ""

            if not current_text:
                current_start = block_start

            current_text = f"{current_text}\n\n{block_text}".strip()
            cursor = offset + block["end"]

            if len(current_text) >= MIN_CHUNK_CHARS and block_text.endswith((".", "?", "!", ":")):
                _append_chunk(
                    chunks,
                    document_id,
                    current_text,
                    current_start,
                    cursor,
                    segment.page,
                    current_heading,
                )
                current_text = ""

        if current_text:
            _append_chunk(
                chunks,
                document_id,
                current_text,
                current_start,
                cursor,
                segment.page,
                current_heading,
            )

    return [asdict(chunk) for chunk in chunks]


def source_segments_from_text(text: str, extension: str) -> list[dict]:
    normalized = _normalize_text(text)
    if not normalized:
        return []

    if extension == "md":
        return [asdict(segment) for segment in _segments_from_markdown(normalized)]

    return [asdict(segment) for segment in _segments_from_plain_text(normalized)]


def source_segments_from_docx(file_bytes: bytes) -> list[dict]:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentIngestionError("python-docx is required to extract DOCX source locations.") from exc

    try:
        document = Document(_bytes_io(file_bytes))
    except (BadZipFile, Exception) as exc:
        raise DocumentIngestionError(f"Malformed DOCX file: {str(exc)}") from exc

    segments: list[SourceSegment] = []
    cursor = 0
    current_heading: str | None = None

    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading"):
            current_heading = text

        end_char = cursor + len(text)
        segments.append(
            SourceSegment(
                text=text,
                heading=current_heading,
                start_char=cursor,
                end_char=end_char,
            )
        )
        cursor = end_char + 2

    if not segments:
        raise DocumentIngestionError("DOCX file did not contain extractable text.")

    return [asdict(segment) for segment in segments]


def source_segments_from_pdf(file_bytes: bytes) -> list[dict]:
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        raise DocumentIngestionError("pypdfium2 is required to extract PDF source locations.") from exc

    try:
        pdf = pdfium.PdfDocument(file_bytes)
    except Exception as exc:
        raise DocumentIngestionError(f"Malformed PDF file: {str(exc)}") from exc

    segments: list[SourceSegment] = []
    for page_index in range(len(pdf)):
        page = pdf[page_index]
        textpage = page.get_textpage()
        page_text = textpage.get_text_range().strip()
        if not page_text:
            continue

        for block in _iter_blocks(_normalize_text(page_text)):
            text = block["text"]
            segments.append(
                SourceSegment(
                    text=text,
                    page=page_index + 1,
                    heading=_heading_from_block(text),
                    start_char=block["start"],
                    end_char=block["end"],
                )
            )

    if not segments:
        raise DocumentIngestionError("PDF file did not contain extractable text.")

    return [asdict(segment) for segment in segments]


def extract_source_segments(file_bytes: bytes, extension: str, fallback_text: str = "") -> list[dict]:
    if extension == "pdf":
        return source_segments_from_pdf(file_bytes)
    if extension == "docx":
        return source_segments_from_docx(file_bytes)
    if extension in {"md", "txt"}:
        text = fallback_text or file_bytes.decode("utf-8", errors="ignore")
        return source_segments_from_text(text, extension)

    raise DocumentIngestionError(
        f"Unsupported file type '{extension}'. "
        f"Only {SOURCE_TRACEABLE_DOCUMENT_LABEL} produce source-traceable document sections."
    )


def _append_chunk(
    chunks: list[DocumentChunk],
    document_id: str,
    text: str,
    start_char: int,
    end_char: int,
    page: int | None,
    heading: str | None,
) -> None:
    index = len(chunks)
    chunks.append(
        DocumentChunk(
            id=deterministic_chunk_id(document_id, index, text),
            document_id=document_id,
            index=index,
            text=text,
            page=page,
            heading=heading,
            start_char=start_char,
            end_char=end_char,
        )
    )


def _coerce_segment(segment: SourceSegment | dict) -> SourceSegment:
    if isinstance(segment, SourceSegment):
        return segment

    return SourceSegment(
        text=str(segment.get("text", "")),
        page=segment.get("page"),
        heading=segment.get("heading"),
        start_char=int(segment.get("start_char") or 0),
        end_char=int(segment.get("end_char") or 0),
    )


def _segments_from_markdown(text: str) -> list[SourceSegment]:
    segments: list[SourceSegment] = []
    current_heading: str | None = None

    for block in _iter_blocks(text):
        block_text = block["text"]
        heading = _heading_from_block(block_text)
        if heading:
            current_heading = heading

        segments.append(
            SourceSegment(
                text=block_text,
                heading=current_heading,
                start_char=block["start"],
                end_char=block["end"],
            )
        )

    return segments


def _segments_from_plain_text(text: str) -> list[SourceSegment]:
    segments = []
    current_heading = None

    for block in _iter_blocks(text):
        block_text = block["text"]
        heading = _heading_from_block(block_text)
        if heading:
            current_heading = heading

        segments.append(
            SourceSegment(
                text=block_text,
                heading=current_heading,
                start_char=block["start"],
                end_char=block["end"],
            )
        )

    return segments


def _iter_blocks(text: str) -> list[dict]:
    blocks = []
    for match in re.finditer(r"\S(?:.*?)(?=\n\s*\n|\Z)", text, re.DOTALL):
        block = match.group(0).strip()
        if block:
            blocks.append({"text": block, "start": match.start(), "end": match.end()})
    return blocks


def _heading_from_block(block: str) -> str | None:
    first_line = block.splitlines()[0].strip()
    if first_line.startswith("#"):
        return first_line.lstrip("#").strip() or None
    if len(first_line) <= 90 and not first_line.endswith((".", "?", "!")):
        return first_line
    return None


def _normalize_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text.replace("\r\n", "\n")).strip()


def clean_source_set_label(label: str | None) -> str:
    return re.sub(r"\s+", " ", str(label or "")).strip()[:160]


def _common_source_root(paths: list[str]) -> str:
    roots = {
        path.split("/", 1)[0]
        for path in paths
        if "/" in path and path.split("/", 1)[0]
    }
    if len(roots) == 1:
        return next(iter(roots))
    return ""


def _stable_source_set_token(value: str) -> str:
    token = re.sub(r"[^A-Za-z0-9_-]+", "-", str(value or "").strip()).strip("-_")
    return token[:80] or "source_set"


def _bytes_io(file_bytes: bytes):
    from io import BytesIO

    return BytesIO(file_bytes)
